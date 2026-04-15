from __future__ import annotations

from pathlib import Path

from PySide6.QtCore import QUrl, Qt
from PySide6.QtGui import QTextCursor, QTextTableFormat, QTextLength
from PySide6.QtWidgets import (
    QStackedWidget,
    QTextBrowser,
    QVBoxLayout,
    QWidget,
)


class DocumentPreviewWidget(QWidget):
    """Виджет просмотра документов с поддержкой режимов DOCX и PDF."""

    def __init__(self, parent: QWidget | None = None, enable_docx_mode: bool = False) -> None:
        super().__init__(parent)
        self._enable_docx_mode = enable_docx_mode
        self._mode = "none"
        self._pdf_doc = None
        self._pdf_view = None
        self._web_view = None
        self._docx_browser = None
        self._fallback_text = None

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        # Stacked widget для переключения между режимами
        self._stack = QStackedWidget()

        # DOCX mode (если включён) — use QTextBrowser for reliable full document rendering
        if self._enable_docx_mode:
            self._docx_web = None
            self._docx_browser = None
            # Use QTextBrowser as primary (more reliable for large documents)
            from PySide6.QtWidgets import QTextBrowser
            self._docx_browser = QTextBrowser()
            self._docx_browser.setReadOnly(True)
            self._docx_browser.setOpenExternalLinks(True)
            self._docx_browser.setLineWrapMode(QTextBrowser.LineWrapMode.WidgetWidth)
            self._docx_browser.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            self._docx_browser.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            self._docx_browser.document().setMaximumBlockCount(0)
            self._docx_browser.setAcceptRichText(True)
            self._stack.addWidget(self._docx_browser)

        # PDF mode
        try:
            from PySide6.QtPdf import QPdfDocument
            from PySide6.QtPdfWidgets import QPdfView

            self._pdf_doc = QPdfDocument(self)
            self._pdf_view = QPdfView(self)
            self._pdf_view.setDocument(self._pdf_doc)
            # Enable multi-page view to show the full document
            self._pdf_view.setPageMode(QPdfView.PageMode.MultiPage)
            self._pdf_widget = QWidget()
            pdf_layout = QVBoxLayout(self._pdf_widget)
            pdf_layout.setContentsMargins(0, 0, 0, 0)
            pdf_layout.addWidget(self._pdf_view)
            self._stack.addWidget(self._pdf_widget)
            self._pdf_index = self._stack.count() - 1
            self._mode = "pdf"
        except Exception:
            self._pdf_index = -1

        # Web fallback
        try:
            from PySide6.QtWebEngineWidgets import QWebEngineView

            self._web_view = QWebEngineView(self)
            self._stack.addWidget(self._web_view)
            if self._mode == "none":
                self._mode = "web"
        except Exception:
            self._web_index = -1

        # Text fallback
        from PySide6.QtWidgets import QTextEdit
        self._fallback_text = QTextEdit(self)
        self._fallback_text.setReadOnly(True)
        self._stack.addWidget(self._fallback_text)

        layout.addWidget(self._stack)

        # Установить начальный режим
        if self._enable_docx_mode:
            self._mode = "docx"
            self._stack.setCurrentIndex(0)  # DOCX по умолчанию
        elif self._mode == "pdf":
            self._stack.setCurrentIndex(self._pdf_index)
        elif self._mode == "web":
            pass  # уже установлен
        else:
            self._mode = "text"
            self._fallback_text.setPlainText("Предпросмотр недоступен: QtPdf и QWebEngine не установлены.")

    def load_docx(self, docx_path: Path) -> None:
        """Загрузить .docx файл и отрендерить в QWebEngineView или QTextBrowser."""
        if not docx_path.exists():
            self._show_error(f"Файл не найден: {docx_path.name}")
            return

        try:
            from docx import Document

            doc = Document(str(docx_path))

            # Collect full HTML content
            html_parts = []
            for element in doc.element.body:
                tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

                if tag == 'p':
                    for para in doc.paragraphs:
                        if para._element == element:
                            html_parts.append(self._render_paragraph_to_html(para))
                            break

                elif tag == 'tbl':
                    for table in doc.tables:
                        if table._tbl == element:
                            html_parts.append(self._render_table_to_html(table))
                            break

            full_html = "<!DOCTYPE html><html><head><meta charset='utf-8'><style>body { font-family: 'Times New Roman', serif; font-size: 12pt; line-height: 1.5; }</style></head><body>" + "".join(html_parts) + "</body></html>"

            # Use QWebEngineView if available, otherwise QTextBrowser
            if self._docx_web is not None:
                self._docx_web.setHtml(full_html)
                self._stack.setCurrentWidget(self._docx_web)
            elif self._docx_browser is not None:
                self._docx_browser.setHtml(full_html)
                self._docx_browser.document().adjustSize()
                self._docx_browser.verticalScrollBar().setValue(0)
                self._docx_browser.update()
                self._stack.setCurrentWidget(self._docx_browser)

        except Exception as exc:
            self._show_error(f"Ошибка рендеринга DOCX: {exc}")

    def _render_paragraph_to_html(self, paragraph) -> str:
        """Convert paragraph to HTML string."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if not paragraph.text.strip() and not paragraph.runs:
            return "<br>"

        alignment = paragraph.alignment
        align_style = ""
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            align_style = 'style="text-align: center;"'
        elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            align_style = 'style="text-align: right;"'
        elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            align_style = 'style="text-align: justify;"'

        html_parts = []
        for run in paragraph.runs:
            style_parts = []
            if run.bold:
                style_parts.append("font-weight: bold;")
            if run.italic:
                style_parts.append("font-style: italic;")
            if run.underline:
                style_parts.append("text-decoration: underline;")
            if run.font.size:
                size_pt = run.font.size.pt
                style_parts.append(f"font-size: {size_pt}pt;")
            if run.font.color and run.font.color.rgb:
                try:
                    color = run.font.color.rgb
                    style_parts.append(f"color: #{color};")
                except Exception:
                    pass

            style_attr = " ".join(style_parts)
            text = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if style_attr:
                html_parts.append(f'<span style="{style_attr}">{text}</span>')
            else:
                html_parts.append(text)

        return f'<p {align_style}>{"".join(html_parts)}</p>'

    def _render_table_to_html(self, table) -> str:
        """Convert table to HTML string."""
        html = '<table border="1" cellpadding="4" cellspacing="0" style="border-collapse: collapse; width: 100%;">'
        for row_idx, row in enumerate(table.rows):
            html += '<tr>'
            for col_idx, cell in enumerate(row.cells):
                cell_html = ""
                for para in cell.paragraphs:
                    if para.runs:
                        for run in para.runs:
                            if run.bold:
                                cell_html += f"<b>{run.text}</b>"
                            elif run.italic:
                                cell_html += f"<i>{run.text}</i>"
                            else:
                                cell_html += run.text
                    else:
                        cell_html += para.text
                    cell_html += "<br>"
                html += f'<td>{cell_html}</td>'
            html += '</tr>'
        html += '</table>'
        return html

    def _render_paragraph(self, cursor: QTextCursor, paragraph) -> None:
        """Отренерить параграф из python-docx в QTextBrowser."""
        from docx.shared import Pt, RGBColor

        if not paragraph.text.strip() and not paragraph.runs:
            cursor.insertHtml("<br>")
            return

        # Определить выравнивание
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        alignment = paragraph.alignment
        align_style = ""
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            align_style = 'style="text-align: center;"'
        elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            align_style = 'style="text-align: right;"'
        elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            align_style = 'style="text-align: justify;"'

        # Собрать HTML
        html_parts = []
        for run in paragraph.runs:
            style_parts = []
            if run.bold:
                style_parts.append("font-weight: bold;")
            if run.italic:
                style_parts.append("font-style: italic;")
            if run.underline:
                style_parts.append("text-decoration: underline;")
            if run.font.size:
                size_pt = run.font.size.pt
                style_parts.append(f"font-size: {size_pt}pt;")
            if run.font.color and run.font.color.rgb:
                try:
                    color = run.font.color.rgb
                    style_parts.append(f"color: #{color};")
                except Exception:
                    pass

            style_attr = " ".join(style_parts)
            text = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if style_attr:
                html_parts.append(f'<span style="{style_attr}">{text}</span>')
            else:
                html_parts.append(text)

        html = f'<p {align_style}>{"".join(html_parts)}</p>'
        cursor.insertHtml(html)
        cursor.insertBlock()

    def _render_table(self, cursor: QTextCursor, table) -> None:
        """Отрендерить таблицу из python-docx в QTextBrowser."""
        from docx.shared import Pt

        rows_count = len(table.rows)
        cols_count = len(table.columns)

        qt_table = cursor.insertTable(rows_count, cols_count)

        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_cursor = qt_table.cellAt(row_idx, col_idx).firstCursorPosition()

                # Собрать текст ячейки
                cell_text = ""
                for para in cell.paragraphs:
                    if para.runs:
                        for run in para.runs:
                            if run.bold:
                                cell_text += f"<b>{run.text}</b>"
                            elif run.italic:
                                cell_text += f"<i>{run.text}</i>"
                            else:
                                cell_text += run.text
                    else:
                        cell_text += para.text
                    cell_text += "<br>"

                if cell_text:
                    cell_cursor.insertHtml(cell_text)

        cursor.movePosition(QTextCursor.AfterRow)
        cursor.insertBlock()

    def load_pdf(self, pdf_path: Path) -> None:
        """Загрузить PDF для предпросмотра."""
        # Check if file is actually a PDF
        if pdf_path.suffix.lower() not in {'.pdf'}:
            # Preview generation failed, show original file or message
            self._show_error(f"Предпросмотр недоступен для {pdf_path.suffix} файлов")
            return

        # Try QPdfView first
        if self._pdf_index >= 0:
            status = self._pdf_doc.load(str(pdf_path.resolve()))
            if status.value != 0:
                self._show_error(f"Не удалось загрузить PDF: {pdf_path}")
            else:
                self._mode = "pdf"
                self._stack.setCurrentIndex(self._pdf_index)
            return

        # Fallback to web view
        if self._web_view is not None:
            self._web_view.setUrl(QUrl.fromLocalFile(str(pdf_path.resolve())))
            self._mode = "web"
            self._stack.setCurrentWidget(self._web_view)
            return

        # Final fallback: show error
        self._show_error(f"Предпросмотр PDF недоступен: {pdf_path}")

    def set_mode(self, mode: str) -> None:
        """Переключить режим просмотра: 'docx' или 'pdf'."""
        if mode == "docx" and self._enable_docx_mode:
            # Use QWebEngineView if available, otherwise QTextBrowser
            if self._docx_web is not None:
                self._mode = "docx"
                self._stack.setCurrentWidget(self._docx_web)
            elif self._docx_browser is not None:
                self._mode = "docx"
                self._stack.setCurrentWidget(self._docx_browser)
        elif mode == "pdf":
            if self._pdf_index >= 0:
                self._mode = "pdf"
                self._stack.setCurrentIndex(self._pdf_index)
            elif self._web_view is not None:
                self._mode = "web"
                self._stack.setCurrentWidget(self._web_view)
            else:
                self._mode = "text"
                self._stack.setCurrentWidget(self._fallback_text)

    def get_current_mode(self) -> str:
        """Получить текущий режим просмотра."""
        return self._mode

    def supports_docx(self) -> bool:
        """Поддерживает ли виджет режим DOCX."""
        return self._enable_docx_mode and (self._docx_web is not None or self._docx_browser is not None)

    def supports_pdf(self) -> bool:
        """Поддерживает ли виджет режим PDF."""
        return self._pdf_index >= 0

    def _show_error(self, text: str) -> None:
        if self._fallback_text is not None:
            self._fallback_text.setPlainText(text)
            if self._mode != "text":
                self._stack.setCurrentWidget(self._fallback_text)

    def show_message(self, text: str) -> None:
        """Показать сообщение в виджете."""
        if self._mode == "text" and self._fallback_text is not None:
            self._fallback_text.setPlainText(text)
            return
        # Fallback: показать в QMessageBox
        try:
            from PySide6.QtWidgets import QMessageBox
            QMessageBox.warning(self, "Предпросмотр", text)
        except Exception:
            pass
