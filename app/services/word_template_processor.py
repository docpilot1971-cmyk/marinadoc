from __future__ import annotations

import copy
import re
from datetime import date
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt

from app.models import ExtractionResult
from app.services.amount_to_words import amount_to_words_ru

_MONTHS_RU_GENITIVE = {
    1: "января",
    2: "февраля",
    3: "марта",
    4: "апреля",
    5: "мая",
    6: "июня",
    7: "июля",
    8: "августа",
    9: "сентября",
    10: "октября",
    11: "ноября",
    12: "декабря",
}


class WordTemplateProcessor:
    def analyze_structure(self, template_path: Path) -> dict[str, object]:
        doc = Document(str(template_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        tables = [{"rows": len(t.rows), "cols": len(t.columns)} for t in doc.tables]
        return {"paragraph_count": len(doc.paragraphs), "non_empty_paragraphs": paragraphs, "tables": tables}

    def render(
        self,
        template_path: Path,
        data: ExtractionResult,
        output_path: Path,
        source_contract_path: Path | None = None,
    ) -> Path:
        from docx.shared import Cm

        doc = Document(str(template_path)) if template_path.exists() else Document()
        placeholders = self._build_placeholder_map(data, source_contract_path)

        # Apply standard page margins
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(1.9)
            section.right_margin = Cm(1.9)

        # Apply paragraph spacing (0 pt before/after, single line spacing)
        for para in doc.paragraphs:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0

        self._insert_contract_table_1_to_1(doc, source_contract_path)
        self._replace_textual_markers(doc, data, placeholders, source_contract_path)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    def _build_placeholder_map(self, data: ExtractionResult, source_contract_path: Path | None) -> dict[str, str]:
        contract_date = data.document.contract_date or self._extract_contract_date_from_contract(source_contract_path)

        # FIX Т1: Дата акта = дата окончания работ, а не today()
        # Если work_end_date_fact раньше contract_date — это ошибка парсера, игнорируем
        raw_work_end = data.period.work_end_date_fact or data.period.work_end_date_plan
        if raw_work_end and contract_date and raw_work_end < contract_date:
            raw_work_end = None  # Parser found a wrong date, ignore it

        act_date = raw_work_end or data.document.act_date or date.today()
        work_end = raw_work_end
        if work_end is None:
            work_end = self._extract_work_end_date_from_contract(source_contract_path)
            # If we extracted work_end from contract, also use it for act_date
            if work_end:
                act_date = work_end

        # Priority: use totals from ExtractionResult (parsed data)
        # Only fallback to contract extraction if result totals are zero/None
        total_value = data.totals.total_with_vat or data.totals.total_without_vat
        if (total_value is None or total_value <= 0) and source_contract_path is not None:
            total_value = self._extract_total_from_contract(source_contract_path)
        if total_value is None:
            total_value = Decimal("0")

        total_digits = self._fmt_money(total_value)
        total_words = amount_to_words_ru(total_value)
        data.totals.total_in_words = total_words

        contract_ref = self._build_contract_ref(data.document.contract_number, contract_date)
        contract_no_short = f"№{data.document.contract_number}" if data.document.contract_number else "№"

        # Signature names in "Фамилия И.О." format
        customer_sig = self._get_signature_name(data.customer.representative_name or data.customer.full_name)
        executor_sig = self._get_signature_name(data.executor.representative_name or data.executor.full_name)

        # IP-specific fields
        executor_name_short = self._extract_short_name(data.executor.full_name or "")

        # FIX Т6: Объект в кавычках
        object_name_raw = data.object_data.object_name or ""
        object_name_quoted = f"«{object_name_raw}»" if object_name_raw and not object_name_raw.startswith("«") else object_name_raw

        # FIX Т8: Должность + ФИО в подписи
        executor_sig_full = f"{data.executor.representative_position or ''} {executor_sig}".strip()
        customer_sig_full = f"{data.customer.representative_position or ''} {customer_sig}".strip()

        return {
            # V1 format markers (for backward compatibility)
            "city_line": f"г. {data.document.document_city or ''} «{act_date.day:02d}» {_MONTHS_RU_GENITIVE[act_date.month]} {act_date.year}г.",
            "customer_intro": (
                f"{data.customer.full_name or ''} в лице {self._to_genitive_fio(data.customer.representative_name) or ''}, "
                f"действующего на основании {data.customer.representative_basis or 'Устава'}, и"
            ),
            "contract_ref": contract_ref,
            "contract_no_short": contract_no_short,
            "object_name": object_name_quoted,
            "object_address": data.object_data.object_address or "",
            "total_digits": total_digits,
            "total_words": total_words,
            "work_end_pretty": self._fmt_ru_date_with_quotes(work_end) if work_end else "",
            "work_end_iso": str(work_end or ""),
            # V2 format markers ({variable} style)
            "city": f"г. {data.document.document_city or ''}",
            "act_day": f"{act_date.day:02d}",
            "act_month": _MONTHS_RU_GENITIVE[act_date.month],
            "act_year": str(act_date.year),
            "contract_number_full": contract_ref,
            "contract_number_short": contract_no_short,
            "customer_full_name": data.customer.full_name or "",
            # FIX Т2: customer_representative — именительный падеж (для строки "Представитель Заказчика: ...")
            # Для "в лице" используется отдельное поле customer_representative_genitive
            "customer_representative": data.customer.representative_name or "",
            "customer_representative_genitive": self._to_genitive_fio(data.customer.representative_name) or "",
            "customer_representative_position": data.customer.representative_position or "",
            "customer_representative_basis": data.customer.representative_basis or "Устава",
            "customer_signature_name": customer_sig,
            "executor_full_name": data.executor.full_name or "",
            "executor_full_name_short": executor_name_short,
            # Executor representative in genitive case for "в лице" construction
            "executor_representative": self._to_genitive_fio(data.executor.representative_name) or "",
            "executor_representative_position": data.executor.representative_position or "",
            # FIX Т3: Fallback на "Устава" для ORG при пустом basis
            "executor_representative_basis": data.executor.representative_basis or "Устава",
            "executor_signature_name": executor_sig,
            "executor_signature_full": executor_sig_full,
            "customer_signature_full": customer_sig_full,
            "work_description": self._build_repair_name(data, source_contract_path=source_contract_path),
            "object_name": object_name_quoted,
            "object_address": data.object_data.object_address or "",
            "work_table": "",  # Table is inserted separately
            "work_end_date_pretty": self._fmt_ru_date_with_quotes(work_end) if work_end else "",
            # IP-specific fields
            "registration_number": self._extract_ip_registration(data.executor.representative_basis or ""),
            "registration_date": self._extract_ip_registration_date(data.executor.representative_basis or ""),
            "tax_inspectorate": self._extract_ip_tax_office(data.executor.representative_basis or "") or "ИФНС",
            "ogrnip": data.executor.ogrnip or "",
            # Full registration string for act template
            "executor_registration_full": data.executor.representative_basis or "",
        }

    @staticmethod
    def _extract_ip_tax_office(basis_text: str) -> str | None:
        """Extract full IFNS name from basis text like '...ИФНС №1 по Белгородской области'."""
        import re
        # Match 'ИФНС' with optional '№X', followed by 'по' and region name
        match = re.search(r"(ИФНС\s*№?\d*\s+по\s+[А-Яа-яЁё\s\-]+(?:области|края|г\.)?)", basis_text, flags=re.IGNORECASE)
        if match:
            return match.group(1).strip()
        # Fallback: just 'ИФНС №X'
        match = re.search(r"(ИФНС\s*№?\d+)", basis_text, flags=re.IGNORECASE)
        if match:
            return match.group(1).strip()
        return None

    @staticmethod
    def _extract_ip_registration_date(basis_text: str) -> str | None:
        """Extract date from 'от DD.MM.YYYY' in registration context."""
        import re
        # Match 'от' followed by date pattern DD.MM.YYYY
        match = re.search(r'от\s+(\d{1,2}\.\d{1,2}\.\d{4})', basis_text)
        if match:
            return match.group(1).strip()
        return None

    @staticmethod
    def _get_signature_name(full_name: str) -> str:
        """Convert 'Иванов Иван Иванович' → 'Иванов И.И.'"""
        if not full_name:
            return ""
        parts = full_name.split()
        if len(parts) >= 3:
            return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
        if len(parts) == 2:
            return f"{parts[0]} {parts[1][0]}."
        return full_name

    @staticmethod
    def _extract_short_name(full_name: str) -> str:
        """Extract short name from full IP name like 'Индивидуальный предприниматель Данков Алексей Николаевич' → 'Данков Алексей Николаевич'."""
        prefixes = ["Индивидуальный предприниматель ", "ИП "]
        for prefix in prefixes:
            if full_name.startswith(prefix):
                return full_name[len(prefix):]
        return full_name

    @staticmethod
    def _extract_ip_registration(basis_text: str) -> str:
        """Extract IP registration number from basis text like 'свидетельства 31 № 001648739 от 14...' → '31 № 001648739'."""
        import re
        match = re.search(r'свидетельства\s+(\d+\s*№\s*\d+)', basis_text)
        if match:
            return match.group(1)
        return ""

    def _replace_textual_markers(
        self,
        doc: Document,
        data: ExtractionResult,
        p: dict[str, str],
        source_contract_path: Path | None,
    ) -> None:
        # Check if template uses {variable} format (V2)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        uses_v2_format = bool(re.search(r'\{(\w+)\}', full_text))

        if uses_v2_format:
            # Use V2 {variable} format replacement
            self._replace_variable_markers(doc, p, source_contract_path)
            return

        # V1 format: #marker# replacement (existing logic)
        for paragraph in doc.paragraphs:
            src = paragraph.text
            text = src

            # Handle city/date special formatting first
            if "#вставить название города формат" in text or "#вставить название города формат г. Город" in text:
                self._format_city_date_paragraph(paragraph, data.document.document_city, data.document.act_date)
                continue

            # Apply ALL marker replacements (not elif - multiple markers can be in one paragraph)
            # For each marker, handle both #marker and #marker# formats
            
            def replace_marker(t: str, marker: str, replacement: str) -> str:
                """Replace marker handling both #marker and #marker# formats."""
                t = t.replace(marker + "#", replacement)  # #marker#
                t = t.replace(marker, replacement)         # #marker
                return t

            # Contract number with full reference
            text = replace_marker(text, "#вставить номер договора из шапки договора №", p["contract_ref"])
            text = replace_marker(text, "#вставить номер договора из шапки договора", p["contract_ref"])
            text = replace_marker(text, "_#вставить номер договора из шапки договора", p["contract_ref"])

            # Customer organization and representative (order matters: specific before general)
            # Handle IP template format where paragraph 4 has full intro with multiple markers
            if "#вставить наименование организации" in text and "#должность и ФИО представителя заказчика" in text:
                # This is the IP template intro line — replace entire customer part
                rep_basis = data.customer.representative_basis or "Устава"
                customer_intro = f"{data.customer.full_name or ''}, в лице {data.customer.representative_name or ''}, действующего на основании {rep_basis}"
                # Find "#вставить наименование организации" and replace everything up to "составили" or "заключили"
                start_marker = "#вставить наименование организации"
                start_idx = text.find(start_marker)
                if start_idx != -1:
                    end_idx = len(text)
                    for em in ["составили", "заключили"]:
                        idx = text.find(em, start_idx)
                        if idx != -1 and idx < end_idx:
                            end_idx = idx
                    before = text[:start_idx]
                    after = text[end_idx:]
                    text = before + customer_intro + " " + after
            elif "#вставить наименование организации, должность и ФИО представителя заказчика" in text:
                text = text.replace("#вставить наименование организации, должность и ФИО представителя заказчика в родительном падеже, действующего на основании ", "")
                text = text.replace("Устава", data.customer.representative_basis or "Устава")
                if data.customer.full_name:
                    text = f"{data.customer.full_name} " + text

            text = replace_marker(text, "#вставить ФИО и данные представителя заказчика", ", ".join(
                [data.customer.representative_name] if data.customer.representative_name else [] +
                [f"действующего на основании {data.customer.representative_basis}"] if data.customer.representative_basis else []
            ) or "")
            text = replace_marker(text, "#должность и ФИО представителя заказчика в родительном падеже", ", ".join(
                [data.customer.representative_name] if data.customer.representative_name else [] +
                [f"действующего на основании {data.customer.representative_basis}"] if data.customer.representative_basis else []
            ) or "")

            # Executor organization and representative (BEFORE generic customer org marker)
            text = replace_marker(text, "#вставить наименование организации Подрядчика", data.executor.full_name or "")
            text = replace_marker(text, "#вставить ФИО представителя подрядчика", data.executor.representative_name or "")
            
            # IP template specific markers
            text = replace_marker(text, "#ФИО представителя подрядчика", data.executor.representative_name or "")
            text = replace_marker(text, "#регистрационный номер", self._extract_ip_registration(data.executor.representative_basis or ""))
            text = replace_marker(text, "#дата выдачи", self._extract_ip_registration_date(data.executor.representative_basis or ""))
            text = replace_marker(text, "#налоговая инспекция", "ИФНС")
            text = replace_marker(text, "#ОГРНИП", data.executor.ogrnip or "")

            # Signature blocks
            # Customer
            cust_pos = data.customer.representative_position or ""
            sig_customer = data.customer.representative_name or data.customer.full_name or ""
            if sig_customer and len(sig_customer.split()) >= 2:
                parts = sig_customer.split()
                sig_customer = f"{parts[0]} {parts[1][0]}." if len(parts) >= 2 else sig_customer
            full_sig_customer = f"{cust_pos} {sig_customer}".strip()
            text = replace_marker(text, "#вставить ФИО и должность представителя заказчика", full_sig_customer)

            # Executor
            exec_pos = data.executor.representative_position or ""
            sig_executor = data.executor.representative_name or data.executor.full_name or ""
            if sig_executor and len(sig_executor.split()) >= 2:
                parts = sig_executor.split()
                sig_executor = f"{parts[0]} {parts[1][0]}." if len(parts) >= 2 else sig_executor
            full_sig_executor = f"{exec_pos} {sig_executor}".strip()
            text = replace_marker(text, "#вставить ФИО и должность представителя подрядчика", full_sig_executor)

            # Generic customer name (AFTER specific markers to avoid partial replacement)
            text = replace_marker(text, "#вставить наименование организации", data.customer.full_name or "")

            # Work details
            if "#вставить наименование ремонтных работ из договора" in text:
                repair_name = self._build_repair_name(data, source_contract_path=source_contract_path)
                text = replace_marker(text, "#вставить наименование ремонтных работ из договора", repair_name)

            # Object name and address — use regex to handle both formats
            obj_pattern = r"#\s*наименование\s+объекта\s+из\s+договora#?"
            if re.search(obj_pattern, text, flags=re.IGNORECASE):
                object_name = p["object_name"]
                if object_name and not object_name.startswith("«"):
                    object_name = f"«{object_name}»"
                text = re.sub(obj_pattern, object_name, text, flags=re.IGNORECASE)

            addr_pattern = r"#\s*(?:вставить\s+)?адрес(?:\s+объекта)?\s+из\s+договора#?"
            if re.search(addr_pattern, text, flags=re.IGNORECASE):
                text = re.sub(addr_pattern, p["object_address"], text, flags=re.IGNORECASE)

            text = replace_marker(text, "#Таблица из договора вставить полностью", "")
            # Handle total — the template may already have "ИТОГО:" before the marker
            total_value_text = f"{p['total_digits']}\n{p['total_words']}"
            text = text.replace("#вставить стоимость в цифрах и прописью рубли, копейки цифрами в расшифровке#", total_value_text)
            text = text.replace("#вставить стоимость в цифрах и прописью рубли, копейки цифрами в расшифровке", total_value_text)
            text = replace_marker(text, "#вставить дату завершения работ из таблицы выполненных работ", p["work_end_pretty"])

            # Special handling for paragraph starting with "В соответствии с договором"
            if "В соответствии с договором строительного подряда" in src:
                repair_name = self._build_repair_name(data, source_contract_path=source_contract_path)
                object_name = p["object_name"]
                if object_name and not object_name.startswith("«"):
                    object_name = f"«{object_name}»"
                text = (
                    f"1. В соответствии с договором строительного подряда {p['contract_ref']} "
                    f"Подрядчик выполнил {repair_name} на объекте {object_name}, "
                    f"находящемся по адресу: {p['object_address']}, а именно:"
                )

            # Clean residual marker prefix if any left
            text = text.replace(" номер договора из шапки договора", "")
            text = text.replace(" наименование ремонтных работ из договора", "")
            text = text.replace(" объекта из договора", "")
            text = text.replace(" адрес объекта из договора", "")
            text = text.replace(" наименование объекта из договора", "")
            text = text.replace(" адрес из договора", "")
            text = re.sub(r"#\S+", "", text)
            if text != src:
                paragraph.text = text
                self._apply_times_new_roman(paragraph)

    def _replace_variable_markers(
        self,
        doc: Document,
        p: dict[str, str],
        source_contract_path: Path | None,
    ) -> None:
        """Replace {variable} markers in V2 templates."""
        for paragraph in doc.paragraphs:
            src = paragraph.text
            text = src

            # Replace all {variable} markers (handle optional trailing space)
            def replace_match(match: re.Match) -> str:
                var_name = match.group(1).strip()
                return p.get(var_name, match.group(0))  # Keep original if not found

            text = re.sub(r'\{([^}]+)\}', replace_match, text)

            # Handle work table insertion
            if "{work_table}" in src:
                # Table is inserted separately via _insert_contract_table_1_to_1
                text = text.replace("{work_table}", "")

            if text != src:
                paragraph.text = text
                self._apply_times_new_roman(paragraph)

    def _insert_contract_table_1_to_1(self, target_doc: Document, source_contract_path: Path | None) -> None:
        source_table_xml = self._extract_primary_contract_table_xml(source_contract_path)
        if source_table_xml is None:
            return

        marker_paragraph = None
        for paragraph in target_doc.paragraphs:
            # Check for both V1 (#marker#) and V2 ({variable}) formats
            if "#Таблица из договора вставить полностью" in paragraph.text or "{work_table}" in paragraph.text:
                marker_paragraph = paragraph
                break
        if marker_paragraph is None:
            return

        marker_paragraph.text = ""
        marker_paragraph._p.addnext(copy.deepcopy(source_table_xml))

        # Add solid borders to the header row (first row) of the inserted table
        self._apply_borders_to_table_header(target_doc)

    @staticmethod
    def _apply_borders_to_table_header(doc: Document) -> None:
        """Apply solid borders to the first row (header) of the first table in the document."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        if not doc.tables:
            return

        table = doc.tables[0]
        header_row = table.rows[0]
        border_val = "single"  # Solid line
        border_sz = "4"  # 4 eighths of a point = 0.5 pt
        border_color = "000000"

        for cell in header_row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                # Remove existing border if any
                existing = tcPr.find(qn(f"w:{border_name}"))
                if existing is not None:
                    tcPr.remove(existing)

                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), border_val)
                border.set(qn("w:sz"), border_sz)
                border.set(qn("w:color"), border_color)
                tcPr.append(border)

    @staticmethod
    def _extract_primary_contract_table_xml(source_contract_path: Path | None):
        if source_contract_path is None or not source_contract_path.exists():
            return None
        source_doc = Document(str(source_contract_path))
        for table in source_doc.tables:
            if not table.rows:
                continue
            header = " ".join(cell.text.lower() for cell in table.rows[0].cells)
            if "№" in header and "наименование" in header and "стоим" in header:
                return table._tbl
        return source_doc.tables[0]._tbl if source_doc.tables else None

    @staticmethod
    def _build_contract_ref(contract_number: str | None, contract_date: date | None) -> str:
        if contract_number and contract_date:
            return f"№{contract_number} от {WordTemplateProcessor._fmt_ru_date_with_quotes(contract_date)}"
        if contract_number:
            return f"№{contract_number}"
        return "№"

    @staticmethod
    def _fmt_ru_date_with_quotes(value: date | None) -> str:
        if value is None:
            return ""
        return f"«{value.day:02d}» {_MONTHS_RU_GENITIVE[value.month]} {value.year}г."

    def _format_city_date_paragraph(self, paragraph, city: str | None, act_date: date | None) -> None:  # noqa: ANN001
        city_value = city or ""
        date_value = self._fmt_ru_date_with_quotes(act_date or date.today())

        paragraph.text = ""
        paragraph.style = paragraph.style
        tab_stops = paragraph.paragraph_format.tab_stops
        tab_stops.clear_all()
        tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT)

        run = paragraph.add_run(f"г. {city_value}\t{date_value}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        self._apply_times_new_roman(paragraph, size_pt=12)

    @staticmethod
    def _apply_times_new_roman(paragraph, size_pt: int | None = None) -> None:  # noqa: ANN001
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            if size_pt is not None:
                run.font.size = Pt(size_pt)

    @staticmethod
    def _fmt_money(value: Decimal | None) -> str:
        if value is None:
            return "0,00"
        quantized = value.quantize(Decimal("0.01"))
        text = f"{quantized:,.2f}".replace(",", " ").replace(".", ",")
        return text

    @staticmethod
    def _build_repair_name(data: ExtractionResult, source_contract_path: Path | None = None) -> str:
        if source_contract_path is not None and source_contract_path.exists():
            source_doc = Document(str(source_contract_path))
            for paragraph in source_doc.paragraphs:
                text = " ".join((paragraph.text or "").split())
                if not text:
                    continue
                match = re.search(
                    r"обязуется\s+выполнить\s+(.+?)\s+на\s+объекте",
                    text,
                    flags=re.IGNORECASE,
                )
                if match:
                    return match.group(1).strip(" ,")
        if data.object_data.object_name:
            return f"на объекте {data.object_data.object_name}"
        return "в соответствии с договором"

    def _extract_total_from_contract(self, source_contract_path: Path | None) -> Decimal | None:
        if source_contract_path is None or not source_contract_path.exists():
            return None
        doc = Document(str(source_contract_path))
        if not doc.tables:
            return None
        table = doc.tables[0]
        # Find explicit "Общая стоимость..." first
        for row in table.rows:
            row_text = " ".join(cell.text.lower() for cell in row.cells)
            if "общая стоимость" in row_text or "итого" in row_text:
                for cell in reversed(row.cells):
                    parsed = self._parse_decimal(cell.text)
                    if parsed is not None:
                        return parsed
        # Fallback: max numeric in last cost column
        max_val: Decimal | None = None
        for row in table.rows:
            if len(row.cells) < 6:
                continue
            parsed = self._parse_decimal(row.cells[5].text)
            if parsed is None:
                continue
            if max_val is None or parsed > max_val:
                max_val = parsed
        return max_val

    def _extract_work_end_date_from_contract(self, source_contract_path: Path | None) -> date | None:
        if source_contract_path is None or not source_contract_path.exists():
            return None
        doc = Document(str(source_contract_path))
        dates: list[date] = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for match in re.finditer(r"\b([0-3]?\d\.[01]?\d\.\d{4})\b", cell.text):
                        try:
                            parsed = date.fromisoformat(
                                "-".join(reversed(match.group(1).split(".")))
                            )
                            dates.append(parsed)
                        except Exception:  # noqa: BLE001
                            continue
        if not dates:
            return None
        return max(dates)

    @staticmethod
    def _to_genitive_fio(fio: str | None) -> str | None:
        if not fio:
            return fio
        parts = fio.split()
        if len(parts) < 2:
            return fio

        def inflect(word: str, idx: int) -> str:
            lower = word.lower()
            is_title = word[:1].isupper()

            # 0 - likely surname, 1 - name, 2 - patronymic
            if idx == 0:
                if lower.endswith("ов") or lower.endswith("ев") or lower.endswith("ин"):
                    out = word + "а"
                elif lower.endswith("ский") or lower.endswith("цкий"):
                    out = word[:-2] + "ого"
                elif lower.endswith("ая"):
                    out = word[:-2] + "ой"
                elif lower.endswith("а"):
                    out = word[:-1] + "ы"
                elif lower.endswith("я"):
                    out = word[:-1] + "и"
                else:
                    out = word
            elif idx == 1:
                if lower.endswith("й"):
                    out = word[:-1] + "я"
                elif lower.endswith("ь"):
                    out = word[:-1] + "я"
                elif lower.endswith("а"):
                    out = word[:-1] + "ы"
                elif lower.endswith("я"):
                    out = word[:-1] + "и"
                else:
                    out = word + "а"
            else:
                if lower.endswith("ич"):
                    out = word + "а"
                elif lower.endswith("на"):
                    out = word[:-1] + "ы"
                else:
                    out = word

            return out if is_title else out.lower()

        converted = [inflect(part, idx) for idx, part in enumerate(parts[:3])]
        if len(parts) > 3:
            converted.extend(parts[3:])
        return " ".join(converted)

    def _extract_contract_date_from_contract(self, source_contract_path: Path | None) -> date | None:
        if source_contract_path is None or not source_contract_path.exists():
            return None
        doc = Document(str(source_contract_path))
        for paragraph in doc.paragraphs[:20]:
            text = " ".join((paragraph.text or "").split())
            if not text:
                continue
            m = re.search(r"«?\s*([0-3]?\d)\s*»?\s+([А-Яа-яЁё]+)\s+(\d{4})\s*г", text)
            if m:
                day, month_name, year = m.groups()
                month_map = {v: k for k, v in _MONTHS_RU_GENITIVE.items()}
                month = month_map.get(month_name.lower())
                if month:
                    try:
                        return date(int(year), month, int(day))
                    except ValueError:
                        pass
            m2 = re.search(r"\b([0-3]?\d\.[01]?\d\.\d{4})\b", text)
            if m2:
                try:
                    return date.fromisoformat("-".join(reversed(m2.group(1).split("."))))
                except Exception:  # noqa: BLE001
                    pass
        return None

    @staticmethod
    def _parse_decimal(value: str) -> Decimal | None:
        cleaned = value.replace("\xa0", "").replace(" ", "").replace(",", ".")
        cleaned = re.sub(r"[^\d.\-]", "", cleaned)
        if not cleaned:
            return None
        try:
            return Decimal(cleaned)
        except Exception:  # noqa: BLE001
            return None
