from __future__ import annotations

import logging
from collections.abc import Iterator
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from app.services.interfaces import IContractReader
from app.services.contract_document import ContractDocument, ParagraphBlock, TableBlock

logger = logging.getLogger(__name__)


class ContractReaderStub(IContractReader):
    def read(self, file_path: Path) -> ContractDocument:
        suffix = file_path.suffix.lower()

        if suffix == ".docx":
            doc = Document(str(file_path))
            parsed = self._read_docx(file_path=file_path, document=doc)
            logger.info(
                "Read DOCX: %s paragraph(s), %s table(s), %s block(s): %s",
                len(parsed.paragraphs),
                len(parsed.tables),
                len(parsed.blocks),
                file_path,
            )
            return parsed

        if suffix == ".doc":
            # TODO: Implement .doc to .docx conversion via pywin32 and then read converted document.
            logger.warning("DOC format is not implemented yet: %s", file_path)
            raise NotImplementedError("DOC support will be added in the next iteration.")

        raise ValueError(f"Unsupported file extension: {suffix}")

    def _read_docx(self, file_path: Path, document: DocxDocument) -> ContractDocument:
        paragraphs: list[str] = []
        tables: list[list[list[str]]] = []
        blocks: list[ParagraphBlock | TableBlock] = []

        for block in self._iter_block_items(document):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if text:
                    paragraphs.append(text)
                    blocks.append(ParagraphBlock(text=text))
                continue

            if isinstance(block, Table):
                rows: list[list[str]] = []
                for row in block.rows:
                    values = [cell.text.strip() for cell in row.cells]
                    if any(values):
                        rows.append(values)
                if rows:
                    tables.append(rows)
                    blocks.append(TableBlock(rows=rows))

        return ContractDocument(file_path=file_path, paragraphs=paragraphs, tables=tables, blocks=blocks)

    @staticmethod
    def _iter_block_items(parent: DocxDocument | _Cell) -> Iterator[Paragraph | Table]:
        if isinstance(parent, DocxDocument):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            return

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)
