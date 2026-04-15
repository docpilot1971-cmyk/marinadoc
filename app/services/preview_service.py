from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

logger = logging.getLogger(__name__)


class PreviewService:
    def __init__(self, preview_dir: Path) -> None:
        self.preview_dir = preview_dir
        self.preview_dir.mkdir(parents=True, exist_ok=True)

    def build_previews(self, generated_files: dict[str, Path]) -> dict[str, Path]:
        previews: dict[str, Path] = {}
        for key, source in generated_files.items():
            if source.suffix.lower() == ".pdf":
                previews[key] = source
                continue
            pdf_target = self.preview_dir / f"{source.stem}.pdf"
            converted = self._convert_to_pdf(source, pdf_target)
            previews[key] = converted if converted else source
        return previews

    def read_document_text(self, source: Path) -> str:
        suffix = source.suffix.lower()
        if suffix in {".docx", ".doc"}:
            return self._read_docx_text(source)
        if suffix in {".xlsx", ".xls"}:
            return self._read_xlsx_text(source)
        return str(source)

    def _convert_to_pdf(self, source: Path, target: Path) -> Path | None:
        suffix = source.suffix.lower()
        try:
            if suffix in {".docx", ".doc"}:
                return self._convert_word(source, target)
            if suffix in {".xlsx", ".xls"}:
                return self._convert_excel(source, target)
        except Exception:  # noqa: BLE001
            logger.exception("Preview conversion failed for %s", source)
            return None
        return None

    @staticmethod
    def _convert_word(source: Path, target: Path) -> Path | None:
        try:
            import pythoncom  # type: ignore[import-not-found]
            import win32com.client  # type: ignore[import-not-found]
            import time
            import gc
        except Exception:  # noqa: BLE001
            return None

        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = None
        try:
            doc = word.Documents.Open(str(source.resolve()), ReadOnly=True)
            time.sleep(0.2)
            doc.ExportAsFixedFormat(str(target.resolve()), 17)
            return target
        except Exception:
            logger.exception("Failed to convert DOCX to PDF: %s", source)
            return None
        finally:
            if doc is not None:
                try:
                    doc.Close(SaveChanges=False)
                except Exception:
                    pass
            if word is not None:
                try:
                    word.Quit(SaveChanges=False)
                except Exception:
                    pass
            doc = None
            word = None
            gc.collect()
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    @staticmethod
    def _convert_excel(source: Path, target: Path) -> Path | None:
        try:
            import pythoncom  # type: ignore[import-not-found]
            import win32com.client  # type: ignore[import-not-found]
            import time
            import gc
        except Exception:  # noqa: BLE001
            return None

        pythoncom.CoInitialize()
        excel = win32com.client.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        wb = None
        try:
            wb = excel.Workbooks.Open(str(source.resolve()), ReadOnly=True)
            time.sleep(0.3)
            wb.Save()
            time.sleep(0.2)
            wb.ExportAsFixedFormat(0, str(target.resolve()))
            return target
        except Exception:
            logger.exception("Failed to convert XLSX to PDF: %s", source)
            return None
        finally:
            if wb is not None:
                try:
                    wb.Close(SaveChanges=False)
                except Exception:
                    pass
            if excel is not None:
                try:
                    excel.Quit()
                except Exception:
                    pass
            wb = None
            excel = None
            gc.collect()
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    @staticmethod
    def _read_docx_text(source: Path) -> str:
        doc = Document(str(source))
        lines: list[str] = []
        for p in doc.paragraphs:
            t = " ".join((p.text or "").split())
            if t:
                lines.append(t)
        for ti, table in enumerate(doc.tables):
            lines.append(f"\n[Table {ti + 1}]")
            for ri, row in enumerate(table.rows[:100]):
                vals = [" ".join((c.text or "").split()) for c in row.cells]
                if any(vals):
                    lines.append(f"R{ri + 1}: " + " | ".join(vals))
        return "\n".join(lines) if lines else f"(Пустой документ: {source.name})"

    @staticmethod
    def _read_xlsx_text(source: Path) -> str:
        wb = load_workbook(source, data_only=False)
        lines: list[str] = []
        for ws in wb.worksheets:
            lines.append(f"[Sheet] {ws.title}")
            shown = 0
            for r in range(1, ws.max_row + 1):
                row_vals: list[str] = []
                for c in range(1, ws.max_column + 1):
                    v = ws.cell(r, c).value
                    if v is not None and str(v).strip() != "":
                        row_vals.append(f"{ws.cell(r, c).coordinate}={v}")
                if row_vals:
                    lines.append(" | ".join(row_vals))
                    shown += 1
                if shown >= 150:
                    lines.append("... preview truncated ...")
                    break
        return "\n".join(lines) if lines else f"(Пустой документ: {source.name})"
