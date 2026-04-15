"""Test parsing and full generation pipeline for DS1 contract."""
import sys
from pathlib import Path
from decimal import Decimal

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from app.core.logging_setup import setup_logging
from app.services.stubs.reader_stub import ContractReaderStub
from app.services.stubs.classifier_stub import ContractTypeClassifierStub
from app.services.stubs.parsers_stub import (
    HeaderParserStub,
    ObjectParserStub,
    PartiesParserStub,
    PeriodParserStub,
    TableParserStub,
    TotalsParserStub,
)
from app.services.excel_template_processor import ExcelTemplateProcessor
from app.services.word_template_processor import WordTemplateProcessor
from app.services.stubs.generators_stub import ActWordGeneratorStub, KS2ExcelGeneratorStub, KS3ExcelGeneratorStub
from app.services.template_loader import TemplateLoader
from app.services.stubs.validator_stub import ExtractionValidatorStub
from app.core.config import AppConfig, PathsConfig, TemplatesConfig
from app.models import ExtractionResult

setup_logging()

# Find the DS contract file
contract_path = PROJECT_ROOT / "templates" / "incoming" / "ДС 1 ПМК-Данков (1).docx"
if not contract_path.exists():
    # Try wildcard
    import glob
    matches = list(PROJECT_ROOT.glob("templates/incoming/*Данков*.docx"))
    if matches:
        contract_path = matches[0]
    else:
        print(f"Contract file not found!")
        sys.exit(1)

print(f"Contract: {contract_path.name}")
print(f"Exists: {contract_path.exists()}")
print()

# Build config
config = AppConfig(
    paths=PathsConfig(
        templates_dir=PROJECT_ROOT / "templates" / "incoming",
        output_dir=PROJECT_ROOT / "output",
        preview_dir=PROJECT_ROOT / "output" / "preview",
    ),
    templates=TemplatesConfig(
        act_word_template="act_org_org_filled_test.docx",
        ks2_template="ks2_test.xlsx",
        ks3_template="ks3_test.xlsx",
    ),
)

# Build services
template_loader = TemplateLoader(config)
word_processor = WordTemplateProcessor()
excel_processor = ExcelTemplateProcessor()
reader = ContractReaderStub()
classifier = ContractTypeClassifierStub()
header_parser = HeaderParserStub()
parties_parser = PartiesParserStub()
object_parser = ObjectParserStub()
period_parser = PeriodParserStub()
table_parser = TableParserStub()
totals_parser = TotalsParserStub()
validator = ExtractionValidatorStub()
act_gen = ActWordGeneratorStub(template_loader, word_processor)
ks2_gen = KS2ExcelGeneratorStub(template_loader, excel_processor)
ks3_gen = KS3ExcelGeneratorStub(template_loader, excel_processor)

# Set source contract for act generator
act_gen.set_source_contract_path(contract_path)

# Full parsing pipeline
print("=" * 80)
print("PARSING")
print("=" * 80)
contract = reader.read(contract_path)
cls = classifier.classify(contract)
print(f"  Type: customer={cls.customer_type}, executor={cls.executor_type}, table={cls.table_grouping_mode}")

doc_data = header_parser.parse(contract)
print(f"  Contract: {doc_data.contract_number} | {doc_data.contract_date} | City: {doc_data.document_city}")

customer, executor = parties_parser.parse(contract)
print(f"  Customer: {customer.full_name}")
print(f"  Customer rep: {customer.representative_name}")
print(f"  Customer basis: {customer.representative_basis}")
print(f"  Executor: {executor.full_name}")
print(f"  Executor rep: {executor.representative_name}")

obj_data = object_parser.parse(contract)
print(f"  Object: {obj_data.object_name[:80] if obj_data.object_name else 'None'}...")
print(f"  Address: {obj_data.object_address[:80] if obj_data.object_address else 'None'}...")

period_data = period_parser.parse(contract)
print(f"  Period: {period_data.work_start_date} — {period_data.work_end_date_fact}")

rows = table_parser.parse(contract, cls.table_grouping_mode)
print(f"  Rows: {len(rows)}")

totals = totals_parser.parse(contract, rows)
print(f"  Total: {totals.total_without_vat} (VAT: {totals.vat_rate}%, {totals.vat_amount})")
print(f"  Total with VAT: {totals.total_with_vat}")

result = ExtractionResult(
    document=doc_data,
    customer=customer,
    executor=executor,
    object_data=obj_data,
    period=period_data,
    rows=rows,
    totals=totals,
)
result.customer.type = cls.customer_type
result.executor.type = cls.executor_type
result = validator.validate(result)

print()
print("=" * 80)
print("GENERATION")
print("=" * 80)

out_dir = PROJECT_ROOT / "output" / "test_ds1"
out_dir.mkdir(parents=True, exist_ok=True)

# Generate act
act_path = out_dir / "act.docx"
try:
    act_gen.generate(result, act_path)
    print(f"  ✅ Act generated: {act_path}")
except Exception as e:
    print(f"  ❌ Act FAILED: {e}")
    import traceback
    traceback.print_exc()

# Generate KS-2
ks2_path = out_dir / "ks2.xlsx"
try:
    ks2_gen.generate(result, ks2_path)
    print(f"  ✅ KS-2 generated: {ks2_path}")
except Exception as e:
    print(f"  ❌ KS-2 FAILED: {e}")
    import traceback
    traceback.print_exc()

# Generate KS-3
ks3_path = out_dir / "ks3.xlsx"
try:
    ks3_gen.generate(result, ks3_path)
    print(f"  ✅ KS-3 generated: {ks3_path}")
except Exception as e:
    print(f"  ❌ KS-3 FAILED: {e}")
    import traceback
    traceback.print_exc()

# Check generated act content
print()
print("=" * 80)
print("ACT CONTENT CHECK")
print("=" * 80)
if act_path.exists():
    from docx import Document
    doc = Document(str(act_path))
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"  {i:2d}: {text[:100]}")
    
    # Check for placeholders
    full_text = "\n".join(p.text for p in doc.paragraphs)
    placeholders = [line for line in full_text.splitlines() if "#" in line]
    if placeholders:
        print(f"\n  ⚠️  UNREPLACED PLACEHOLDERS ({len(placeholders)}):")
        for ph in placeholders[:5]:
            print(f"    {ph[:80]}")
    else:
        print(f"\n  ✅ All placeholders replaced")
    
    # Check for key data
    checks = {
        "Contract number": doc_data.contract_number or "❌ NOT FOUND",
        "Customer name": customer.full_name[:50] if customer.full_name else "❌ NOT FOUND",
        "Customer rep": customer.representative_name or "❌ NOT FOUND",
        "City": doc_data.document_city or "❌ NOT FOUND",
        "Object": (obj_data.object_name or "❌ NOT FOUND")[:50],
    }
    print()
    for key, val in checks.items():
        status = "✅" if "❌" not in str(val) else "⚠️ "
        print(f"  {status} {key}: {val}")
